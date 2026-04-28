#!/usr/bin/env python3
"""
scripts/office/unpack.py
Desempaqueta un .docx y describe su estructura interna.

Uso:
    python3 scripts/office/unpack.py <ruta.docx> [--format=json|text] [--outline-only]

Salida (stdout):
    JSON estructurado (por defecto) o texto legible (--format=text).
    Código de salida 0 = OK, 1 = error.

Información extraída:
    - Propiedades del documento (título, autor, fecha de creación/modificación)
    - Esquema de capítulos (Heading 1/2/3 con texto y nivel)
    - Estilos utilizados (únicos, frecuencia)
    - Conteo de párrafos, tablas, imágenes
    - Estimación de páginas
    - Lista de imágenes embebidas (nombre de archivo interno)
    - Lista de tablas (filas × columnas, primera fila como encabezado)
"""

import sys
import os
import json
import zipfile
import argparse
import xml.etree.ElementTree as ET
from pathlib import Path
from collections import Counter, defaultdict
from datetime import datetime

# python-docx es preferible para parsear estilos de forma fiable
try:
    import docx
    from docx import Document
    from docx.oxml.ns import qn
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False


# ---------------------------------------------------------------------------
# Namespaces OOXML
# ---------------------------------------------------------------------------
NS = {
    'w':   'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'r':   'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
    'cp':  'http://schemas.openxmlformats.org/package/2006/metadata/core-properties',
    'dc':  'http://purl.org/dc/elements/1.1/',
    'dcterms': 'http://purl.org/dc/terms/',
    'a':   'http://schemas.openxmlformats.org/drawingml/2006/main',
    'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
    'wp':  'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
}

HEADING_STYLES = {
    'heading 1': 1, 'heading 2': 2, 'heading 3': 3,
    'heading 4': 4, 'heading 5': 5, 'heading 6': 6,
    'título 1': 1, 'título 2': 2, 'título 3': 3,
    'encabezado 1': 1, 'encabezado 2': 2, 'encabezado 3': 3,
}


# ---------------------------------------------------------------------------
# Extraer propiedades del documento (core-properties)
# ---------------------------------------------------------------------------
def _extract_core_props(zf: zipfile.ZipFile) -> dict:
    props = {}
    try:
        with zf.open('docProps/core.xml') as f:
            root = ET.parse(f).getroot()
        for tag, key in [
            ('{http://purl.org/dc/elements/1.1/}title',           'title'),
            ('{http://purl.org/dc/elements/1.1/}creator',         'author'),
            ('{http://purl.org/dc/elements/1.1/}description',     'description'),
            ('{http://schemas.openxmlformats.org/package/2006/metadata/core-properties}lastModifiedBy', 'last_modified_by'),
            ('{http://purl.org/dc/terms/}created',                'created'),
            ('{http://purl.org/dc/terms/}modified',               'modified'),
            ('{http://schemas.openxmlformats.org/package/2006/metadata/core-properties}revision', 'revision'),
        ]:
            el = root.find(tag)
            if el is not None and el.text:
                props[key] = el.text.strip()
    except KeyError:
        pass  # core.xml no existe
    return props


# ---------------------------------------------------------------------------
# Extraer estructura via python-docx (preferido)
# ---------------------------------------------------------------------------
def _analyze_with_python_docx(path: Path) -> dict:
    doc = Document(str(path))

    outline = []
    style_counter = Counter()
    para_count = 0
    word_count = 0

    for para in doc.paragraphs:
        para_count += 1
        style_name = para.style.name if para.style else 'Normal'
        style_counter[style_name] += 1
        text = para.text.strip()
        words = len(text.split()) if text else 0
        word_count += words

        style_lower = style_name.lower()
        level = HEADING_STYLES.get(style_lower)
        if level is not None and text:
            outline.append({'level': level, 'style': style_name, 'text': text})

    # Tablas
    tables_info = []
    for i, table in enumerate(doc.tables):
        rows = len(table.rows)
        cols = len(table.columns) if table.rows else 0
        header = []
        if table.rows:
            try:
                header = [cell.text.strip() for cell in table.rows[0].cells]
            except Exception:
                header = []
        tables_info.append({
            'index': i,
            'rows': rows,
            'cols': cols,
            'header_row': header,
        })

    # Secciones
    sections_info = []
    for i, sec in enumerate(doc.sections):
        sections_info.append({
            'index': i,
            'orientation': 'landscape' if sec.orientation == 1 else 'portrait',
            'page_width_cm':  round(sec.page_width.cm, 1)  if sec.page_width  else None,
            'page_height_cm': round(sec.page_height.cm, 1) if sec.page_height else None,
        })

    return {
        'paragraphs':  para_count,
        'word_count':  word_count,
        'outline':     outline,
        'styles_used': [{'style': s, 'count': c} for s, c in style_counter.most_common()],
        'tables':      tables_info,
        'sections':    sections_info,
    }


# ---------------------------------------------------------------------------
# Fallback: análisis directo de XML (sin python-docx)
# ---------------------------------------------------------------------------
def _analyze_with_xml(zf: zipfile.ZipFile) -> dict:
    try:
        with zf.open('word/document.xml') as f:
            root = ET.parse(f).getroot()
    except KeyError:
        raise ValueError("El archivo no contiene word/document.xml — ¿es un .docx válido?")

    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    outline = []
    style_counter = Counter()
    para_count = 0
    word_count = 0

    for para in root.iter(f'{{{W}}}p'):
        para_count += 1
        ppr = para.find(f'{{{W}}}pPr')
        style_name = 'Normal'
        if ppr is not None:
            pstyle = ppr.find(f'{{{W}}}pStyle')
            if pstyle is not None:
                style_name = pstyle.get(f'{{{W}}}val', 'Normal')
        style_counter[style_name] += 1

        text = ''.join(t.text or '' for t in para.iter(f'{{{W}}}t')).strip()
        word_count += len(text.split()) if text else 0

        style_lower = style_name.lower()
        level = HEADING_STYLES.get(style_lower)
        if level is None:
            # Intentar por prefijo numérico tipo "heading1", "Heading1"
            for k, v in HEADING_STYLES.items():
                if style_lower.startswith(k.replace(' ', '')):
                    level = v
                    break
        if level is not None and text:
            outline.append({'level': level, 'style': style_name, 'text': text})

    tables_info = []
    for i, tbl in enumerate(root.iter(f'{{{W}}}tbl')):
        rows = list(tbl.iter(f'{{{W}}}tr'))
        cols = max((len(list(r.iter(f'{{{W}}}tc'))) for r in rows), default=0)
        header = []
        if rows:
            header = [''.join(t.text or '' for t in rows[0].iter(f'{{{W}}}t')).strip()
                      for cell in rows[0].iter(f'{{{W}}}tc')]
        tables_info.append({'index': i, 'rows': len(rows), 'cols': cols, 'header_row': header})

    return {
        'paragraphs': para_count,
        'word_count': word_count,
        'outline':    outline,
        'styles_used': [{'style': s, 'count': c} for s, c in style_counter.most_common()],
        'tables':     tables_info,
        'sections':   [],
    }


# ---------------------------------------------------------------------------
# Imágenes embebidas
# ---------------------------------------------------------------------------
def _list_images(zf: zipfile.ZipFile) -> list:
    images = []
    for name in zf.namelist():
        lower = name.lower()
        if lower.startswith('word/media/') and any(
            lower.endswith(ext) for ext in ('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.emf', '.wmf', '.svg')
        ):
            info = zf.getinfo(name)
            images.append({
                'filename': os.path.basename(name),
                'path_in_zip': name,
                'size_bytes': info.file_size,
            })
    return images


# ---------------------------------------------------------------------------
# Estimar páginas (heurística: ~350 palabras por página)
# ---------------------------------------------------------------------------
def _estimate_pages(word_count: int, table_count: int) -> int:
    words_equiv = word_count + table_count * 150
    return max(1, round(words_equiv / 350))


# ---------------------------------------------------------------------------
# Salida texto legible
# ---------------------------------------------------------------------------
def _format_text(result: dict) -> str:
    lines = []
    lines.append('=' * 60)
    lines.append(f"DOCUMENTO: {result['file']}")
    lines.append('=' * 60)

    props = result.get('properties', {})
    if props:
        lines.append('\n[PROPIEDADES]')
        for k, v in props.items():
            lines.append(f'  {k}: {v}')

    lines.append('\n[ESTADÍSTICAS]')
    stats = result.get('stats', {})
    lines.append(f"  Párrafos:         {stats.get('paragraphs', '?')}")
    lines.append(f"  Palabras (aprox): {stats.get('word_count', '?')}")
    lines.append(f"  Páginas (est.):   {stats.get('pages_estimated', '?')}")
    lines.append(f"  Tablas:           {stats.get('table_count', '?')}")
    lines.append(f"  Imágenes:         {stats.get('image_count', '?')}")
    lines.append(f"  Secciones:        {stats.get('section_count', '?')}")

    outline = result.get('outline', [])
    if outline:
        lines.append('\n[ESQUEMA / ÍNDICE]')
        for item in outline:
            indent = '  ' * (item['level'] - 1)
            lines.append(f"  {indent}{'#' * item['level']} {item['text']}")

    tables = result.get('tables', [])
    if tables:
        lines.append('\n[TABLAS]')
        for t in tables:
            lines.append(f"  Tabla {t['index'] + 1}: {t['rows']} filas × {t['cols']} cols")
            if t.get('header_row'):
                header_str = ' | '.join(str(h)[:30] for h in t['header_row'][:6])
                lines.append(f"    Encabezado: {header_str}")

    images = result.get('images', [])
    if images:
        lines.append('\n[IMÁGENES]')
        for img in images:
            lines.append(f"  {img['filename']}  ({img['size_bytes']:,} bytes)")

    styles = result.get('styles_used', [])
    if styles:
        lines.append('\n[ESTILOS (top 15)]')
        for s in styles[:15]:
            lines.append(f"  {s['count']:4d}× {s['style']}")

    lines.append('')
    return '\n'.join(lines)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
def main():
    parser = argparse.ArgumentParser(
        description='Desempaqueta un .docx y describe su estructura interna.',
        formatter_class=argparse.RawTextHelpFormatter,
    )
    parser.add_argument('docx_path', help='Ruta al archivo .docx')
    parser.add_argument(
        '--format', choices=['json', 'text'], default='json',
        help='Formato de salida: json (por defecto) o text'
    )
    parser.add_argument(
        '--outline-only', action='store_true',
        help='Mostrar solo el esquema de capítulos (reduce output)'
    )
    args = parser.parse_args()

    path = Path(args.docx_path)

    # Validaciones
    if not path.exists():
        print(json.dumps({'error': f'Archivo no encontrado: {path}'}), file=sys.stderr)
        sys.exit(1)
    if path.suffix.lower() != '.docx':
        print(json.dumps({'error': f'El archivo no es .docx: {path.name}'}), file=sys.stderr)
        sys.exit(1)
    if not zipfile.is_zipfile(path):
        print(json.dumps({'error': f'El archivo no es un ZIP válido (¿.docx corrupto?): {path.name}'}), file=sys.stderr)
        sys.exit(1)

    try:
        with zipfile.ZipFile(path, 'r') as zf:
            props   = _extract_core_props(zf)
            images  = _list_images(zf)

            if HAS_PYTHON_DOCX:
                analysis = _analyze_with_python_docx(path)
            else:
                analysis = _analyze_with_xml(zf)

        pages_est = _estimate_pages(analysis['word_count'], len(analysis['tables']))

        result = {
            'file': str(path),
            'parser': 'python-docx' if HAS_PYTHON_DOCX else 'xml-stdlib',
            'properties': props,
            'stats': {
                'paragraphs':      analysis['paragraphs'],
                'word_count':      analysis['word_count'],
                'pages_estimated': pages_est,
                'table_count':     len(analysis['tables']),
                'image_count':     len(images),
                'section_count':   len(analysis.get('sections', [])),
            },
            'outline':     analysis['outline'],
            'tables':      analysis['tables'],
            'images':      images,
            'styles_used': analysis['styles_used'],
            'sections':    analysis.get('sections', []),
        }

        if args.outline_only:
            result = {
                'file':    result['file'],
                'outline': result['outline'],
                'stats':   result['stats'],
            }

        if args.format == 'text':
            print(_format_text(result))
        else:
            print(json.dumps(result, ensure_ascii=False, indent=2))

    except Exception as exc:
        print(json.dumps({'error': str(exc), 'file': str(path)}), file=sys.stderr)
        sys.exit(1)


if __name__ == '__main__':
    main()
