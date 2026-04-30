#!/usr/bin/env python3
from pathlib import Path
import shutil
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.shared import RGBColor, Pt

ROOT = Path(__file__).resolve().parents[1]
BASE3 = ROOT / 'DOCS - ANEJOS' / '3.- Estudio Geotecnico' / 'Anexo 3 - Estudio Geotecnico.docx'
BASE5 = ROOT / 'DOCS - ANEJOS' / '5.- Dimensionamiento del firme' / 'Anexo 5 - Dimensionamiento del Firme.docx'
REV3 = ROOT / 'DOCS - ANEJOS' / '3.- Estudio Geotecnico' / 'Anexo 3 - Estudio Geotecnico_REV.docx'
REV5 = ROOT / 'DOCS - ANEJOS' / '5.- Dimensionamiento del firme' / 'Anexo 5 - Dimensionamiento del Firme_REV.docx'

COLOR_HEADER = '366092'
COLOR_ALT = 'D9EAF7'


def safe_copy(src, dst):
    dst.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(src, dst)


def set_cell_shading(cell, fill_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tc_pr.append(shd)


def style_cell_text(cell, bold=False, white=False, size=10):
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = bold
            r.font.size = Pt(size)
            if white:
                r.font.color.rgb = RGBColor(255, 255, 255)
            else:
                r.font.color.rgb = RGBColor(0, 0, 0)


def style_table(table):
    try:
        table.style = 'Table Grid'
    except Exception:
        pass

    if len(table.rows) == 0:
        return

    # Header row
    for c in table.rows[0].cells:
        set_cell_shading(c, COLOR_HEADER)
        style_cell_text(c, bold=True, white=True, size=10)

    # Body rows
    for i, row in enumerate(table.rows[1:], start=1):
        for c in row.cells:
            if i % 2 == 0:
                set_cell_shading(c, COLOR_ALT)
            style_cell_text(c, bold=False, white=False, size=10)


def add_par_after(par, text, style='Normal'):
    new_p = OxmlElement('w:p')
    par._p.addnext(new_p)
    out = Paragraph(new_p, par._parent)
    if style:
        out.style = style
    if text:
        out.add_run(text)
    return out


def add_table_after(doc, par, rows):
    table = doc.add_table(rows=0, cols=len(rows[0]))
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    style_table(table)

    tbl = table._tbl
    tbl.getparent().remove(tbl)
    par._p.addnext(tbl)
    new_p = OxmlElement('w:p')
    tbl.addnext(new_p)
    return Paragraph(new_p, par._parent)


def find_heading(doc, token):
    token = token.lower()
    for p in doc.paragraphs:
        if p.style and 'Heading' in p.style.name and token in (p.text or '').lower():
            return p
    return None


def build_anexo3():
    safe_copy(BASE3, REV3)
    doc = Document(REV3)

    h = find_heading(doc, '1. objeto')
    if h:
        p = add_par_after(h, 'El presente anejo recopila y justifica los criterios geotecnicos aplicables al nuevo viario de la Ampliacion Plaza Mayor, con especial atencion a la explanada y a su compatibilidad con el firme de referencia del ambito (Aparcamiento 1, Aparcamiento 2 y Eje Norte).')
        add_par_after(p, 'La redaccion se basa en documentacion geotecnica historica y documentacion de proyecto anterior del propio ambito, sin campana geotecnica nueva especifica para este encargo. Por tanto, la aplicacion final de parametros debe verificarse en la fase de obra.')

    h = find_heading(doc, '3. trabajos realizados')
    if h:
        p = add_par_after(h, 'Se ha realizado lectura integra de la documentacion geotecnica y de firme del proyecto anterior, extraccion de datos de trazabilidad y contraste de coherencia para su aplicacion a la ampliacion.')
        p = add_par_after(p, 'Tabla 1. Documentacion analizada')
        add_table_after(doc, p, [
            ['Documento', 'Tipo', 'Uso tecnico en este anejo'],
            ['Proyecto Plaza Mayor incluidas las modificaciones (PDF)', 'Proyecto anterior', 'Fuente principal para seccion de firme, criterios de explanada y zonificacion geotecnica'],
            ['6523 Plaza Mayor Shopping (PDF)', 'Estudio geotecnico previo', 'Contexto general de materiales y reconocimientos del entorno'],
            ['6863 bis Manzana 8.1 SUNP BM 3 (PDF)', 'Sondeos geotecnicos', 'Contraste local de litologia y nivel freatico en la manzana de referencia'],
            ['Anejo 3 del proyecto de referencia', 'Guia estructural', 'Modelo de orden semantico, bloques y nivel de desarrollo tecnico'],
        ])

    h = find_heading(doc, '5. estudio de la explanada')
    if h:
        p = add_par_after(h, 'La solucion de explanada del nuevo viario se define por continuidad tecnica con el proyecto anterior, evitando cambios bruscos de rigidez entre plataformas y asegurando un apoyo homogeneo bajo el paquete granular y bituminoso.')
        p = add_par_after(p, 'Tabla 2. Explanada mejorada adoptada')
        add_table_after(doc, p, [
            ['Elemento', 'Valor adoptado'],
            ['Espesor', '25 cm'],
            ['Material', 'Suelo seleccionado tipo S-2'],
            ['Procedencia', 'Prestamo'],
            ['CBR', '> 20'],
            ['Compactacion', '95% Proctor Modificado'],
        ])

        p = add_par_after(p, 'Tabla 3. Controles minimos de ejecucion geotecnica')
        add_table_after(doc, p, [
            ['Control', 'Criterio minimo'],
            ['Cajeo y saneo', 'Retirada de materiales inadecuados y blandones detectados'],
            ['Regularizacion de apoyo', 'Plano homogeneo antes de la capa de mejora'],
            ['Compactacion por tongadas', 'Control de humedad y densidad en cada tongada'],
            ['Control de espesor', 'Verificacion del espesor total de 25 cm en la mejora'],
            ['Control documental', 'Registro de ensayos y trazabilidad por zonas de obra'],
        ])

    last = doc.paragraphs[-1]
    h6 = add_par_after(last, '6. CONCLUSIONES', 'Heading 2')
    p = add_par_after(h6, '1) El ambito presenta condicionantes geotecnicos propios de materiales aluviales heterogeneos y exige control de ejecucion de caja y apoyo.')
    p = add_par_after(p, '2) Se adopta una explanada mejorada de 25 cm con suelo S-2 de prestamo, CBR > 20 y compactacion al 95% Proctor Modificado como criterio de coronacion del cimiento del firme.')
    p = add_par_after(p, '3) La solucion es coherente con el firme de referencia del proyecto anterior y debe validarse en obra en funcion de condiciones reales de humedad y capacidad portante.')

    h7 = add_par_after(p, '7. DATOS PENDIENTES', 'Heading 2')
    p = add_par_after(h7, 'Tabla 4. Datos pendientes para cierre geotecnico')
    add_table_after(doc, p, [
        ['Dato', 'Estado'],
        ['Categoria de trafico IMD/IMDp', '[PENDIENTE: confirmar dato de proyecto si se exige cierre normativo por categoria]'],
        ['Inventario final de ensayos historicos', '[PENDIENTE: consolidar apendices para archivo definitivo]'],
        ['Validacion de arranque de obra', '[PENDIENTE: inspeccion geotecnica de caja por Direccion Facultativa]'],
    ])

    doc.save(REV3)


def build_anexo5():
    safe_copy(BASE5, REV5)
    doc = Document(REV5)

    h = find_heading(doc, '1. objeto')
    if h:
        p = add_par_after(h, 'El presente anejo define y justifica la solucion de firme del nuevo viario de la Ampliacion Plaza Mayor, adoptando como criterio principal la continuidad estructural y funcional con la seccion de referencia de Aparcamiento 1, Aparcamiento 2 y Eje Norte.')
        add_par_after(p, 'El objetivo es evitar cambios bruscos de rigidez, asegurar transiciones constructivas correctas con zonas ya ejecutadas/proyectadas y mantener coherencia tecnica entre explanada, capas granulares y capas bituminosas.')

    h = find_heading(doc, '2. metodologia')
    if h:
        p = add_par_after(h, 'El dimensionamiento se plantea con metodologia de continuidad con el proyecto anterior del ambito, complementada con criterios normativos de referencia para materiales y ejecucion.')
        p = add_par_after(p, 'Se toma como base el PG-3 para suelos seleccionados, zahorras, riegos y mezclas bituminosas. La Norma 6.1-IC se emplea como marco tecnico de contraste cuando existan datos suficientes de trafico y explanada. La Norma 6.3-IC se considera solo de forma complementaria para encuentros con firme existente cuando aplique.')
        p = add_par_after(p, 'Tabla 1. Marco normativo y criterio de aplicacion')
        add_table_after(doc, p, [
            ['Referencia', 'Aplicacion en este anejo'],
            ['PG-3', 'Base de especificaciones de materiales, ejecucion y control de capas'],
            ['Norma 6.1-IC', 'Referencia tecnica de contraste para clasificacion por trafico y secciones'],
            ['Norma 6.3-IC', 'Referencia complementaria en transiciones o compatibilidades con firmes existentes'],
            ['Proyecto anterior Plaza Mayor', 'Fuente principal de la seccion adoptada por continuidad'],
        ])
        add_par_after(p, 'La nomenclatura de mezclas D-12 y G-25 se mantiene expresamente para conservar trazabilidad documental con el proyecto anterior, sin perjuicio de adaptar denominaciones en fase de pliego/mediciones, si procede.')

    h = find_heading(doc, '3.- categoria de trafico')
    if h:
        p = add_par_after(h, 'No se dispone en la documentacion analizada de IMD/IMDp especifica del nuevo viario que permita fijar con trazabilidad una categoria de trafico pesado propia para este anejo.')
        p = add_par_after(p, '[PENDIENTE: confirmar categoria de trafico o intensidad de vehiculos pesados del nuevo viario].')
        add_par_after(p, 'En consecuencia, no se fuerza una categoria T sin base documental y se justifica la solucion por continuidad con el firme de referencia del ambito, criterio tecnicamente coherente para un viario interior de urbanizacion con conexion directa a plataformas ya definidas.')

    h = find_heading(doc, '4.- explanada')
    if h:
        p = add_par_after(h, 'Como coronacion/mejora del cimiento del firme se adopta una capa de explanada mejorada de 25 cm de espesor con suelo seleccionado S-2 de prestamo, CBR > 20 y compactacion al 95% Proctor Modificado, para homogeneizar apoyo y controlar deformabilidad local.')
        p = add_par_after(p, 'Tabla 2. Explanada mejorada adoptada')
        add_table_after(doc, p, [
            ['Elemento', 'Caracteristicas'],
            ['Explanada mejorada', '25 cm'],
            ['Material', 'Suelo seleccionado tipo S-2'],
            ['Procedencia', 'Prestamo'],
            ['CBR', '> 20'],
            ['Compactacion', '95% Proctor Modificado'],
        ])

    h = find_heading(doc, '5.1. calzada')
    if h:
        p = add_par_after(h, 'Se adopta para calzada la seccion de referencia del proyecto anterior, manteniendo composicion y espesores para garantizar continuidad funcional, geometrica y estructural con Aparcamiento 1, Aparcamiento 2 y Eje Norte.')
        p = add_par_after(p, 'Tabla 3. Seccion de firme de referencia')
        add_table_after(doc, p, [
            ['Capa', 'Espesor'],
            ['Mezcla bituminosa D-12', '5 cm'],
            ['Riego de adherencia', '-'],
            ['Mezcla bituminosa G-25', '9 cm'],
            ['Riego de imprimacion', '-'],
            ['Zahorra artificial', '25 cm'],
            ['Zahorra natural', '25 cm'],
            ['Explanada compactada', '-'],
        ])

        p = add_par_after(p, 'Tabla 4. Justificacion tecnica de capas y funcion estructural')
        add_table_after(doc, p, [
            ['Capa/elemento', 'Funcion tecnica principal'],
            ['D-12 (5 cm)', 'Capa superior de rodadura y regularidad superficial'],
            ['G-25 (9 cm)', 'Capa estructural bituminosa de reparto de cargas'],
            ['Riegos', 'Asegurar adherencia/ligazon entre capas'],
            ['Zahorra artificial 25 cm', 'Base granular de alta capacidad y uniformidad'],
            ['Zahorra natural 25 cm', 'Subbase de transicion y distribucion de tensiones'],
            ['Explanada mejorada S-2', 'Homogeneizacion del apoyo y control de asentamientos diferenciales'],
        ])

    h = find_heading(doc, '5.2. aceras')
    if h:
        add_par_after(h, 'Se mantiene criterio de pavimento peatonal sobre apoyo regularizado y compactado, con especial atencion al drenaje, juntas y encuentros con bordillo/calzada. [PENDIENTE: cierre de seccion definitiva de acera segun planos y pliego vigentes].')

    h = find_heading(doc, '5.3. carril bici')
    if h:
        add_par_after(h, '[PENDIENTE: confirmar alcance real de carril bici en la ampliacion y definir, en su caso, seccion especifica compatible con el resto del ambito].')

    last = doc.paragraphs[-1]
    h6 = add_par_after(last, '6. CONDICIONES DE EJECUCION Y CONTROL', 'Heading 2')
    p = add_par_after(h6, 'Tabla 5. Controles minimos de obra recomendados')
    add_table_after(doc, p, [
        ['Control', 'Criterio minimo'],
        ['Cajeo y saneo', 'Retirada de materiales inadecuados y blandones'],
        ['Control de humedad', 'No ejecutar compactacion fuera de rango admisible'],
        ['Compactacion por tongadas', 'Verificacion de densidad y humedad por capa'],
        ['Espesores', 'Comprobacion sistematica de espesores de capas granulares y bituminosas'],
        ['Riegos bituminosos', 'Limpieza previa y dosificacion/ejecucion correcta'],
        ['Encuentros firme nuevo/existente', 'Tratamiento de transiciones para evitar discontinuidades de rigidez'],
    ])

    h7 = add_par_after(h6, '7. CONCLUSIONES', 'Heading 2')
    p = add_par_after(h7, '1) Se adopta la seccion de referencia del proyecto anterior: 5 cm D-12, riego de adherencia, 9 cm G-25, riego de imprimacion, 25 cm zahorra artificial, 25 cm zahorra natural y explanada compactada.')
    p = add_par_after(p, '2) Se incorpora explanada mejorada de 25 cm con suelo S-2 de prestamo, CBR > 20 y compactacion al 95% Proctor Modificado como elemento de coronacion del cimiento.')
    p = add_par_after(p, '3) La solucion garantiza continuidad funcional, geometrica y estructural con las zonas de referencia del ambito y minimiza riesgo de cambios bruscos de comportamiento.')
    p = add_par_after(p, '4) La categoria de trafico propia del nuevo viario queda expresamente pendiente de confirmacion documental, sin forzar clasificaciones no sustentadas.')

    h8 = add_par_after(p, '8. DATOS PENDIENTES', 'Heading 2')
    p = add_par_after(h8, 'Tabla 6. Datos pendientes')
    add_table_after(doc, p, [
        ['Dato', 'Estado'],
        ['IMD/IMDp del nuevo viario', '[PENDIENTE: confirmar dato de proyecto]'],
        ['Seccion definitiva de aceras', '[PENDIENTE: validar con planos/pliego]'],
        ['Seccion de carril bici', '[PENDIENTE: confirmar si aplica en alcance final]'],
    ])

    doc.save(REV5)


def main():
    build_anexo3()
    build_anexo5()
    print(REV3)
    print(REV5)


if __name__ == '__main__':
    main()
