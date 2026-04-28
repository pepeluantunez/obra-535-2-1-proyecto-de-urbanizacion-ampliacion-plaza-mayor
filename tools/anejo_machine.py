#!/usr/bin/env python3
"""
anejo_machine.py — Máquina de anejos para 535.2.1 Plaza Mayor
==============================================================
Punto de entrada único para gestionar, auditar y sincronizar los anejos del
Proyecto de Urbanización - Ampliación Plaza Mayor (expediente 535.2.1).

Comandos:
  status          — Tabla de estado de todos los anejos
  check <n>       — Auditoría detallada de un anejo (número o slug)
  sync-tables <n> — Sincronizar tablas del donor Guadalmar al anejo de Plaza Mayor
  refresh <n>     — Re-aplicar apertura y datos de proyecto al anejo
  list-donors     — Listar disponibilidad de donors en Guadalmar

Uso:
  python3 tools/anejo_machine.py status
  python3 tools/anejo_machine.py check 14
  python3 tools/anejo_machine.py sync-tables 14 [--dry-run]
  python3 tools/anejo_machine.py list-donors

Reglas:
  - Nunca escribe directamente datos técnicos: solo estructura y textos de apertura
  - sync-tables crea backup antes de modificar
  - Los datos numéricos provienen de los donors, nunca del LLM
"""

import argparse
import copy
import json
import os
import re
import shutil
import sys
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor
    import lxml.etree as etree
except ImportError:
    print("ERROR: Instalar python-docx — pip install python-docx --break-system-packages")
    sys.exit(1)

# ---------------------------------------------------------------------------
# Rutas
# ---------------------------------------------------------------------------

SCRIPT_DIR = Path(__file__).resolve().parent
PROJECT_ROOT = SCRIPT_DIR.parent
CONFIG_DIR = PROJECT_ROOT / "CONFIG"
ANEJOS_DIR = PROJECT_ROOT / "DOCS - ANEJOS"
CONTROL_DIR = PROJECT_ROOT / "CONTROL"

APERTURA_JSON = CONFIG_DIR / "apertura_anejos_plaza_mayor.json"
PROFILES_JSON = CONFIG_DIR / "annex_template_profiles.json"
PROYECTO_JSON = CONFIG_DIR / "proyecto.template.json"

def _find_guadalmar_root() -> Path:
    """Busca el directorio raiz de Guadalmar en posibles rutas de montaje."""
    pou = "535.2.2 Mejora Carretera Guadalmar/POU 2026"
    candidates = [
        PROJECT_ROOT.parent / "MEJORA CARRETERA GUADALMAR/PROYECTO 535/535.2" / pou,
        PROJECT_ROOT.parent / "Projects/MEJORA CARRETERA GUADALMAR/PROYECTO 535/535.2" / pou,
        Path("/sessions/busy-awesome-ptolemy/mnt/POU 2026"),
        Path("/sessions/busy-awesome-ptolemy/mnt/Projects/MEJORA CARRETERA GUADALMAR/PROYECTO 535/535.2") / pou,
    ]
    for c in candidates:
        if c.exists():
            return c
    return candidates[0]

GUADALMAR_ROOT = _find_guadalmar_root()

# ---------------------------------------------------------------------------
# Helpers de carga
# ---------------------------------------------------------------------------

def load_json(path: Path) -> dict:
    with open(path, encoding="utf-8") as f:
        return json.load(f)


def load_apertura() -> dict:
    data = load_json(APERTURA_JSON)
    return {a["number"]: a for a in data["annexes"]}


def load_profiles() -> dict:
    data = load_json(PROFILES_JSON)
    return {p["number"]: p for p in data["profiles"]}


def load_proyecto() -> dict:
    return load_json(PROYECTO_JSON)


# ---------------------------------------------------------------------------
# Mapeo de carpetas de anejos
# ---------------------------------------------------------------------------

FOLDER_MAP = {
    1: "1.- Reportaje Fotografico",
    2: "2.- Cartografia y Topografia",
    3: "3.- Estudio Geotecnico",
    4: "4.- Trazado, Replanteo y Mediciones Auxiliares",
    5: "5.- Dimensionamiento del Firme",
    6: "6.- Red de Agua Potable",
    7: "7.- Red de Saneamiento - Pluviales",
    8: "8.- Red de Saneamiento - Fecales",
    9: "9.- Red de Media Tension",
    10: "10.- Red de Baja Tension",
    11: "11.- Red de Alumbrado",
    12: "12.- Accesibilidad",
    13: "13.- Estudio de Gestion de Residuos",
    14: "14.- Control de Calidad",
    15: "15.- Plan de Obra",
    16: "16.- Comunicaciones con Companias Suministradoras",
    17: "17.- Seguridad y Salud",
    18: "18.- Telecomunicaciones",
}

DOCX_NAME_MAP = {
    1: "Anexo 1 - Reportaje Fotografico.docx",
    2: "Anexo 2 - Cartografia y Topografia.docx",
    3: "Anexo 3 - Estudio Geotecnico.docx",
    4: "Anexo 4 - Trazado, Replanteo y Mediciones Auxiliares.docx",
    5: "Anexo 5 - Dimensionamiento del Firme.docx",
    6: "Anexo 6 - Red de Agua Potable.docx",
    7: "Anexo 7 - Red de Saneamiento - Pluviales.docx",
    8: "Anexo 8 - Red de Saneamiento - Fecales.docx",
    9: "Anexo 9 - Red de Media Tension.docx",
    10: "Anexo 10 - Red de Baja Tension.docx",
    11: "Anexo 11 - Red de Alumbrado.docx",
    12: "Anexo 12 - Accesibilidad.docx",
    13: "Anexo 13 - Estudio de Gestion de Residuos.docx",
    14: "Anexo 14 - Control de Calidad.docx",
    15: "Anexo 15 - Plan de Obra.docx",
    16: "Anexo 16 - Comunicaciones con Companias Suministradoras.docx",
    17: "Anexo 17 - Estudio de Seguridad y Salud.docx",
    18: "Anexo 18 - Telecomunicaciones.docx",
}


# Anejos fuera de alcance según CLAUDE.md (eléctrico + telecomunicaciones)
OUT_OF_SCOPE = {9, 10, 11, 18}


def anejo_docx_path(n: int) -> Path | None:
    folder = FOLDER_MAP.get(n)
    docx = DOCX_NAME_MAP.get(n)
    if not folder or not docx:
        return None
    candidates = list((ANEJOS_DIR / folder).glob("*.docx")) if (ANEJOS_DIR / folder).exists() else []
    # Prefer canonical name, fall back to first found
    canonical = ANEJOS_DIR / folder / docx
    if canonical.exists():
        return canonical
    for c in candidates:
        if "backup" not in c.name.lower() and "_bak" not in c.name.lower():
            return c
    return None


def donor_docx_path(n: int, profiles: dict) -> Path | None:
    profile = profiles.get(n)
    if not profile:
        return None
    rel = profile.get("donor_docx", "")
    if not rel:
        return None
    # Paths in JSON are relative to the project root (parent of tools/)
    candidate = (PROJECT_ROOT / rel).resolve()
    if candidate.exists():
        return candidate
    # Try relative to GUADALMAR_ROOT
    fname = Path(rel).name
    for root, dirs, files in os.walk(GUADALMAR_ROOT / "DOCS" / "Documentos de Trabajo"):
        for f in files:
            if f == fname:
                return Path(root) / f
    return None


# ---------------------------------------------------------------------------
# Análisis de un DOCX
# ---------------------------------------------------------------------------

PLACEHOLDER_PATTERNS = [
    r"Tabla \d+\. Descripcion",
    r"\[RELLENAR\]",
    r"\[PENDIENTE\]",
    r"\bRELLENAR\b",
    r"\bPENDIENTE\b",
    r"\[.*?\]",
    r"PRECIO PENDIENTE",
]
PLACEHOLDER_RE = re.compile("|".join(PLACEHOLDER_PATTERNS), re.IGNORECASE)

GUADALMAR_MARKERS = [
    "guadalmar", "535.2.2", "mejora de la carretera", "carretera de guadalmar",
]


def audit_docx(path: Path) -> dict:
    """Devuelve un dict con estadísticas del docx."""
    if not path or not path.exists():
        return {"exists": False}
    try:
        doc = Document(str(path))
    except Exception as e:
        return {"exists": True, "error": str(e)}

    paras = [p.text.strip() for p in doc.paragraphs]
    non_empty = [p for p in paras if p]
    tables = doc.tables

    placeholders = []
    guadalmar_refs = []
    for p in non_empty:
        if PLACEHOLDER_RE.search(p):
            placeholders.append(p[:80])
        for m in GUADALMAR_MARKERS:
            if m in p.lower():
                guadalmar_refs.append(p[:80])
                break

    # Check table placeholder cells
    for t in tables:
        for row in t.rows:
            for cell in row.cells:
                txt = cell.text.strip()
                if PLACEHOLDER_RE.search(txt):
                    placeholders.append(f"[TABLE] {txt[:60]}")
                for m in GUADALMAR_MARKERS:
                    if m in txt.lower():
                        guadalmar_refs.append(f"[TABLE] {txt[:60]}")
                        break

    size_kb = path.stat().st_size // 1024

    return {
        "exists": True,
        "path": str(path),
        "paragraphs": len(non_empty),
        "tables": len(tables),
        "placeholders": placeholders,
        "guadalmar_refs": guadalmar_refs,
        "size_kb": size_kb,
    }


def health_score(audit: dict) -> str:
    """Semáforo: VERDE / AMARILLO / ROJO / VACÍO / ERROR."""
    if not audit.get("exists"):
        return "VACÍO"
    if "error" in audit:
        return "ERROR"
    if audit.get("guadalmar_refs"):
        return "⚠ GUADALMAR"
    n_ph = len(audit.get("placeholders", []))
    n_para = audit.get("paragraphs", 0)
    if n_para < 5:
        return "VACÍO"
    if n_ph == 0 and n_para > 20:
        return "✓ OK"
    if n_ph <= 3:
        return "~ REVISAR"
    return "✗ PLACEHOLDERS"


# ---------------------------------------------------------------------------
# COMANDO: status
# ---------------------------------------------------------------------------

def cmd_status(args):
    apertura = load_apertura()
    profiles = load_profiles()

    print(f"\n{'='*90}")
    print(f"ESTADO ANEJOS — 535.2.1 Ampliación Plaza Mayor   [{datetime.now().strftime('%Y-%m-%d')}]")
    print(f"{'='*90}")
    fmt = "{:>3}  {:<42} {:>6} {:>6} {:>4}  {:>8}  {:<16} {}"
    print(fmt.format("N.", "Título", "Párrs", "Tablas", "KB", "Placeh.", "Estado", "Donor"))
    print("-" * 90)

    for n in range(1, 19):
        title = apertura.get(n, {}).get("title", f"ANEJO {n}")[:41]
        docx = anejo_docx_path(n)
        audit = audit_docx(docx)
        status = health_score(audit)
        has_profile = n in profiles
        donor_ok = "✓ " if (has_profile and donor_docx_path(n, profiles)) else ("○ " if has_profile else "—")

        if audit.get("exists"):
            print(fmt.format(
                n, title,
                audit.get("paragraphs", 0),
                audit.get("tables", 0),
                audit.get("size_kb", 0),
                len(audit.get("placeholders", [])),
                status,
                donor_ok,
            ))
        elif n in OUT_OF_SCOPE:
            print(fmt.format(n, title, "—", "—", "—", "—", "FUERA ALCANCE", donor_ok))
        else:
            print(fmt.format(n, title, "—", "—", "—", "—", "FALTA DOCX", donor_ok))

    print()
    print("Donor: ✓ =disponible | ○ =perfil existe sin donor local | — =sin perfil")
    print("Estado: ✓ OK | ~ REVISAR | ✗ PLACEHOLDERS | ⚠ GUADALMAR | VACÍO | FALTA DOCX | FUERA ALCANCE")
    print()


# ---------------------------------------------------------------------------
# COMANDO: check
# ---------------------------------------------------------------------------

def cmd_check(args):
    n = int(args.anejo)
    profiles = load_profiles()
    apertura = load_apertura()

    anejo_info = apertura.get(n, {})
    title = anejo_info.get("title", f"ANEJO {n}")
    print(f"\n{'='*70}")
    print(f"CHECK ANEJO {n}: {title}")
    print(f"{'='*70}")

    docx = anejo_docx_path(n)
    audit = audit_docx(docx)

    if not audit.get("exists"):
        print(f"❌  DOCX no encontrado en: {ANEJOS_DIR / FOLDER_MAP.get(n, '?')}")
    else:
        print(f"Archivo : {audit['path']}")
        print(f"Tamaño  : {audit['size_kb']} KB")
        print(f"Párrafos: {audit['paragraphs']}")
        print(f"Tablas  : {audit['tables']}")
        print(f"Estado  : {health_score(audit)}")

        if audit.get("placeholders"):
            print(f"\nPlaceholders encontrados ({len(audit['placeholders'])}):")
            for p in audit["placeholders"][:15]:
                print(f"  • {p}")
        else:
            print("\nSin placeholders detectados ✓")

        if audit.get("guadalmar_refs"):
            print(f"\n⚠  Referencias a Guadalmar ({len(audit['guadalmar_refs'])}):")
            for r in audit["guadalmar_refs"][:10]:
                print(f"  >> {r}")
        else:
            print("Sin referencias a Guadalmar ✓")

    # Donor
    profile = profiles.get(n)
    if profile:
        donor = donor_docx_path(n, profiles)
        print(f"\nDonor configurado: {Path(profile['donor_docx']).name}")
        if donor and donor.exists():
            d_audit = audit_docx(donor)
            print(f"Donor disponible : ✓  ({d_audit.get('paragraphs')} párrs, {d_audit.get('tables')} tablas, {d_audit.get('size_kb')} KB)")
        else:
            print(f"Donor disponible : ✗  NO ENCONTRADO")
            print(f"  Buscado en: {profile['donor_docx']}")

        if profile.get("variable_inputs"):
            print(f"\nVariables a rellenar ({len(profile['variable_inputs'])}):")
            for v in profile["variable_inputs"]:
                print(f"  {v}")
    else:
        print(f"\nSin donor configurado (anejo {n} se genera desde plantilla maestra)")

    # Apertura
    if anejo_info.get("object_text"):
        print(f"\nTexto de objeto (apertura_anejos_plaza_mayor.json):")
        print(f"  {anejo_info['object_text'][:200]}...")
    print()


# ---------------------------------------------------------------------------
# COMANDO: list-donors
# ---------------------------------------------------------------------------

def cmd_list_donors(args):
    profiles = load_profiles()
    print(f"\n{'='*80}")
    print(f"DONORS GUADALMAR → 535.2.1   [{datetime.now().strftime('%Y-%m-%d')}]")
    print(f"{'='*80}")
    fmt = "{:>3}  {:<35} {:<8} {}"
    print(fmt.format("N.", "Donor file", "Estado", "Ruta local"))
    print("-" * 80)

    for n in sorted(profiles.keys()):
        profile = profiles[n]
        fname = Path(profile["donor_docx"]).name
        donor = donor_docx_path(n, profiles)
        if donor and donor.exists():
            status = "✓ OK"
            ruta = str(donor)[-60:]
        else:
            status = "✗ FALTA"
            ruta = profile["donor_docx"][-60:]
        print(fmt.format(n, fname[:34], status, ruta))

    anejos_sin_donor = [n for n in range(1, 19) if n not in profiles]
    print(f"\nAnejos sin donor (generan desde plantilla): {anejos_sin_donor}")
    print()


# ---------------------------------------------------------------------------
# COMANDO: sync-tables
# ---------------------------------------------------------------------------

def _backup_docx(path: Path) -> Path:
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    bak = path.parent / f"_bak_{ts}_{path.name}"
    shutil.copy2(str(path), str(bak))
    return bak


def _copy_table_xml(src_table, dst_doc) -> object:
    """Copia la definición XML de una tabla del donor al documento destino."""
    src_xml = copy.deepcopy(src_table._tbl)
    return src_xml


def cmd_sync_tables(args):
    n = int(args.anejo)
    dry_run = getattr(args, "dry_run", False)

    profiles = load_profiles()
    apertura = load_apertura()

    profile = profiles.get(n)
    if not profile:
        print(f"ERROR: Anejo {n} no tiene donor configurado en annex_template_profiles.json")
        sys.exit(1)

    target_path = anejo_docx_path(n)
    if not target_path or not target_path.exists():
        print(f"ERROR: Anejo {n} de Plaza Mayor no encontrado en DOCS - ANEJOS/")
        sys.exit(1)

    donor_path = donor_docx_path(n, profiles)
    if not donor_path or not donor_path.exists():
        print(f"ERROR: Donor no encontrado. Ruta configurada: {profile['donor_docx']}")
        sys.exit(1)

    print(f"\nsync-tables  Anejo {n}  {'[DRY-RUN]' if dry_run else ''}")
    print(f"  Target : {target_path.name}")
    print(f"  Donor  : {donor_path.name}")

    # Cargar ambos
    target_doc = Document(str(target_path))
    donor_doc = Document(str(donor_path))

    target_tables = target_doc.tables
    donor_tables = donor_doc.tables

    print(f"  Target tables: {len(target_tables)}")
    print(f"  Donor  tables: {len(donor_tables)}")

    if not donor_tables:
        print("  INFO: El donor no tiene tablas. Nada que sincronizar.")
        return

    # Auditar target para detectar placeholders en tablas
    placeholder_tables = []
    for i, t in enumerate(target_tables):
        all_text = " ".join(c.text for row in t.rows for c in row.cells)
        if PLACEHOLDER_RE.search(all_text):
            placeholder_tables.append(i)

    print(f"  Tablas placeholder en target: {placeholder_tables}")

    if dry_run:
        print(f"\n  [DRY-RUN] Se reemplazarían {min(len(placeholder_tables), len(donor_tables))} tablas.")
        print("  Ejecutar sin --dry-run para aplicar cambios.")
        return

    # Backup obligatorio
    bak = _backup_docx(target_path)
    print(f"  Backup: {bak.name}")

    # Reemplazar tablas placeholder con las del donor (en orden)
    replaced = 0
    for ph_idx, donor_table in zip(placeholder_tables, donor_tables):
        target_tbl = target_tables[ph_idx]._tbl
        new_tbl = copy.deepcopy(donor_table._tbl)
        target_tbl.getparent().replace(target_tbl, new_tbl)
        replaced += 1
        print(f"  → Tabla {ph_idx + 1} reemplazada desde donor")

    # Si hay más donor tables sin placeholder correspondiente, informar
    if len(donor_tables) > len(placeholder_tables):
        extra = len(donor_tables) - len(placeholder_tables)
        print(f"  INFO: {extra} tablas donor sin placeholder correspondiente — no insertadas")
        print(f"  Para insertar tablas adicionales, usar 'append-donor-tables' (pendiente)")

    target_doc.save(str(target_path))
    print(f"  ✓ Guardado: {target_path.name}")
    print(f"  Tablas reemplazadas: {replaced}")

    # Post-check
    check_doc = Document(str(target_path))
    ph_after = 0
    for t in check_doc.tables:
        all_text = " ".join(c.text for row in t.rows for c in row.cells)
        if PLACEHOLDER_RE.search(all_text):
            ph_after += 1
    if ph_after == 0:
        print("  ✓ Sin placeholders en tablas tras sync.")
    else:
        print(f"  ⚠  Quedan {ph_after} tablas con posibles placeholders.")

    # Verificar mojibake
    _check_mojibake(target_path)
    print()


def _check_mojibake(path: Path):
    """Verifica ausencia de mojibake en el DOCX resultante."""
    MOJIBAKE = ["\xc3\x83", "\xc3\x82", "â€", "COMPROBACIÃ\x93N", "URBANIZACIÃ\x93N"]
    try:
        doc = Document(str(path))
        found = []
        for p in doc.paragraphs:
            for token in MOJIBAKE:
                if token in p.text:
                    found.append(f"párrafo: {p.text[:60]}")
                    break
        if found:
            print(f"  ⚠  MOJIBAKE DETECTADO ({len(found)} casos):")
            for f in found[:5]:
                print(f"     {f}")
        else:
            print("  ✓ Sin mojibake detectado.")
    except Exception as e:
        print(f"  ⚠  No se pudo verificar mojibake: {e}")


# ---------------------------------------------------------------------------
# COMANDO: refresh
# ---------------------------------------------------------------------------

def cmd_refresh(args):
    n = int(args.anejo)
    apertura = load_apertura()
    proyecto = load_proyecto()

    anejo_info = apertura.get(n, {})
    if not anejo_info:
        print(f"ERROR: Anejo {n} no encontrado en apertura_anejos_plaza_mayor.json")
        sys.exit(1)

    target_path = anejo_docx_path(n)
    if not target_path or not target_path.exists():
        print(f"ERROR: Anejo {n} de Plaza Mayor no encontrado")
        sys.exit(1)

    dry_run = getattr(args, "dry_run", False)
    print(f"\nrefresh  Anejo {n}  {'[DRY-RUN]' if dry_run else ''}")
    print(f"  Archivo: {target_path.name}")

    doc = Document(str(target_path))

    # Reemplazos de identidad del proyecto
    REPLACEMENTS = {
        "MEJORA DE LA CARRETERA DE GUADALMAR": "AMPLIACIÓN PLAZA MAYOR",
        "Mejora Carretera Guadalmar": "Ampliación Plaza Mayor",
        "535.2.2": "535.2.1",
        "Guadalmar": "Plaza Mayor",
        proyecto.get("project_name", "PROYECTO"): proyecto.get("project_name", "PROYECTO"),
    }

    changes = []
    for para in doc.paragraphs:
        for old, new in REPLACEMENTS.items():
            if old in para.text and old != new:
                for run in para.runs:
                    if old in run.text:
                        run.text = run.text.replace(old, new)
                        changes.append(f"párrafo: '{old}' → '{new}'")

    # Reemplazar texto de objeto si hay un párrafo "OBJETO" vacío
    object_text = anejo_info.get("object_text", "")
    if object_text:
        for i, para in enumerate(doc.paragraphs):
            if re.match(r"^(El presente anejo tiene por objeto|OBJETO)", para.text.strip(), re.IGNORECASE):
                if len(para.text.strip()) < 30:  # párrafo vacío o stub
                    if not dry_run:
                        para.runs[0].text = object_text if para.runs else ""
                    changes.append(f"objeto: insertado texto de apertura ({len(object_text)} chars)")
                    break

    print(f"  Cambios detectados: {len(changes)}")
    for c in changes[:10]:
        print(f"  • {c}")

    if dry_run:
        print("  [DRY-RUN] No se guardan cambios.")
        return

    if not changes:
        print("  INFO: Sin cambios necesarios.")
        return

    bak = _backup_docx(target_path)
    print(f"  Backup: {bak.name}")
    doc.save(str(target_path))
    _check_mojibake(target_path)
    print(f"  ✓ Guardado: {target_path.name}")
    print()


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------


# ---------------------------------------------------------------------------
# COMANDO: fix-captions
# ---------------------------------------------------------------------------

def cmd_fix_captions(args):
    n = int(args.anejo)
    dry_run = getattr(args, "dry_run", False)

    target_path = anejo_docx_path(n)
    if not target_path or not target_path.exists():
        print(f"ERROR: Anejo {n} no encontrado")
        sys.exit(1)

    print(f"\nfix-captions  Anejo {n}  {'[DRY-RUN]' if dry_run else ''}")
    print(f"  Archivo: {target_path.name}")

    doc = Document(str(target_path))
    body = doc.element.body

    # Mapear: para cada párrafo "Tabla N. Descripcion", encontrar la sección
    # precedente (heading) más cercana y construir un caption real.
    CAPTION_RE = re.compile(r"^Tabla\s+\d+\.\s+Descripcion$", re.IGNORECASE)
    NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    children = list(body)
    changes = []
    last_heading = ""

    table_counter = {}  # by section prefix -> count

    for i, child in enumerate(children):
        tag = child.tag.split("}")[-1]
        if tag == "p":
            # Get text
            text = "".join(
                r.text or "" for r in child.iter(f"{{{NS}}}t")
            ).strip()
            # Track last heading (style Heading or numbered section)
            if re.match(r"^\d+[\.\d]*\s+", text) or "CAPÍTULO" in text.upper():
                last_heading = text[:80]
            # Detect caption placeholder
            if CAPTION_RE.match(text):
                section_key = re.sub(r"[^A-Za-z0-9]", "_", last_heading[:20])
                table_counter[section_key] = table_counter.get(section_key, 0) + 1
                # Build new caption
                new_caption = f"Tabla. {last_heading[:60]}" if last_heading else f"Tabla."
                changes.append({
                    "element": child,
                    "old": text,
                    "new": new_caption,
                    "index": i,
                })

    print(f"  Captions placeholder encontrados: {len(changes)}")
    for c in changes:
        print(f"  • [{c['old']}] → [{c['new']}]")

    if dry_run:
        print("\n  [DRY-RUN] No se guardan cambios.")
        return

    if not changes:
        print("  INFO: Sin captions placeholder.")
        return

    bak = _backup_docx(target_path)
    print(f"  Backup: {bak.name}")

    for c in changes:
        elem = c["element"]
        # Clear runs and set text
        for r in list(elem.findall(f"{{{NS}}}r")):
            elem.remove(r)
        # Create a new run
        new_r = etree.SubElement(elem, f"{{{NS}}}r")
        new_t = etree.SubElement(new_r, f"{{{NS}}}t")
        new_t.text = c["new"]
        new_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")

    doc.save(str(target_path))
    _check_mojibake(target_path)
    print(f"  ✓ {len(changes)} captions actualizados → {target_path.name}")
    print()

def main():
    parser = argparse.ArgumentParser(
        description="anejo_machine.py — Máquina de anejos 535.2.1 Plaza Mayor",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__,
    )
    sub = parser.add_subparsers(dest="command")

    sub.add_parser("status", help="Tabla de estado de todos los anejos")

    p_check = sub.add_parser("check", help="Auditoría detallada de un anejo")
    p_check.add_argument("anejo", help="Número de anejo (1-18)")

    p_sync = sub.add_parser("sync-tables", help="Sincronizar tablas del donor al anejo")
    p_sync.add_argument("anejo", help="Número de anejo (1-18)")
    p_sync.add_argument("--dry-run", action="store_true", help="Simular sin escribir")

    sub.add_parser("list-donors", help="Listar disponibilidad de donors en Guadalmar")

    p_refresh = sub.add_parser("refresh", help="Re-aplicar apertura y datos de proyecto")
    p_refresh.add_argument("anejo", help="Número de anejo (1-18)")
    p_refresh.add_argument("--dry-run", action="store_true", help="Simular sin escribir")

    p_fc = sub.add_parser("fix-captions", help="Reemplazar captions placeholder 'Tabla N. Descripcion'")
    p_fc.add_argument("anejo", help="Número de anejo (1-18)")
    p_fc.add_argument("--dry-run", action="store_true")

    args = parser.parse_args()

    if args.command == "status":
        cmd_status(args)
    elif args.command == "check":
        cmd_check(args)
    elif args.command == "sync-tables":
        cmd_sync_tables(args)
    elif args.command == "list-donors":
        cmd_list_donors(args)
    elif args.command == "refresh":
        cmd_refresh(args)
    elif args.command == "fix-captions":
        cmd_fix_captions(args)
    else:
        parser.print_help()
        sys.exit(1)


if __name__ == "__main__":
    main()


