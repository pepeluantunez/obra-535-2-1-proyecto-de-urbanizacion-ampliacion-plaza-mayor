#!/usr/bin/env python3
"""
apply_project_style.py
Estandariza el estilo visual de Excel y Word del expediente POU.

Aplica una identidad visual coherente de proyecto a todos los archivos:
- Excel (openpyxl): cabeceras con color corporativo, filas alternas, anchos,
  congelacion de paneles, formato numerico espanol, print setup.
- Word (python-docx): tablas con cabecera consistente, bordes uniformes,
  fuente homogenea en todo el documento.

Los DATOS no se tocan. Solo formato.

Uso:
    python3 tools/apply_project_style.py [opciones]

Opciones:
    --excel-dir DIR    Carpeta donde buscar xlsx (defecto: DOCS - ANEJOS/)
    --word-dir DIR     Carpeta donde buscar docx (defecto: DOCS - ANEJOS/)
    --dry-run          Analizar sin modificar
    --excel-only       Solo Excel
    --word-only        Solo Word
    --file RUTA        Procesar un archivo concreto
    --backup           Crear copia _ORIG antes de modificar (defecto: si)

PALETA DEL PROYECTO 535.2.x — extraida de las 3 plantillas maestras Word:
    Fuente:        Montserrat (identidad del proyecto — Word y Excel)
    Header bg:     #366092 (azul corporativo — unanime en PLANTILLA_MAESTRA_*.docx)
    Header text:   #FFFFFF (blanco)
    Alt row:       #D9EAF7 (azul muy palido — 6084 usos en tablas del proyecto)
    Subtotal bg:   #D9E2F3 (azul suave — usado en Excel GR y SyS)
    Dark accent:   #0E2841 (dk2 del tema — titulos y bordes fuertes)
    Link/accent:   #0070C0 (azul brillante — texto secundario)
    Grey neutral:  #F2F2F2 (gris claro — filas sin datos / separadores)
    Fuente size:   10pt Excel / 11pt Word
"""

import os
import sys
import shutil
import argparse
import re
from pathlib import Path
from copy import copy
from datetime import datetime

# ── Dependencias ──────────────────────────────────────────────────────────────
try:
    import openpyxl
    from openpyxl.styles import (
        PatternFill, Font, Alignment, Border, Side, numbers
    )
    from openpyxl.utils import get_column_letter
    HAVE_OPENPYXL = True
except ImportError:
    HAVE_OPENPYXL = False

try:
    import docx
    from docx.shared import Pt, RGBColor, Cm
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    HAVE_DOCX = True
except ImportError:
    HAVE_DOCX = False


# ── Paleta del proyecto ───────────────────────────────────────────────────────
# Extraida de PLANTILLA_MAESTRA_ANEJOS/MEMORIA/ESS.docx (3 plantillas unanimes)
class Palette:
    HEADER_BG     = '366092'   # azul corporativo (cabeceras Word y Excel)
    HEADER_FG     = 'FFFFFF'   # blanco
    ALT_ROW       = 'D9EAF7'   # azul muy palido (fila alterna — 6084 usos en tablas)
    SUBTOTAL_BG   = 'D9E2F3'   # azul suave (subtotales / agrupaciones)
    DARK_ACCENT   = '0E2841'   # azul muy oscuro (tema dk2 — bordes fuertes)
    BORDER_COLOR  = '366092'   # igual que cabecera para coherencia
    GREY_NEUTRAL  = 'F2F2F2'   # gris claro (separadores / filas sin datos)
    WHITE         = 'FFFFFF'
    BLACK         = '000000'
    FONT_NAME     = 'Montserrat'   # fuente del proyecto (31.555 usos en anejos)
    FONT_SIZE_XL  = 10             # Excel
    FONT_SIZE_DOC = 11             # Word


# ── Helpers colores ───────────────────────────────────────────────────────────
def _fill(hex_color):
    return PatternFill('solid', fgColor=hex_color)

def _font(bold=False, color=Palette.BLACK, size=Palette.FONT_SIZE_XL):
    return Font(name=Palette.FONT_NAME, bold=bold, color=color, size=size)

def _border_side(style='thin'):
    return Side(style=style, color=Palette.BORDER_COLOR)

def _border(style='thin'):
    s = _border_side(style)
    return Border(left=s, right=s, top=s, bottom=s)

def _align(h='left', v='center', wrap=False):
    return Alignment(horizontal=h, vertical=v, wrap_text=wrap)


# ══════════════════════════════════════════════════════════════════════════════
# EXCEL
# ══════════════════════════════════════════════════════════════════════════════

def _is_header_row(ws, row_idx, col_min, col_max):
    """Heuristica: primera fila no vacia de la hoja."""
    for col in range(col_min, col_max + 1):
        cell = ws.cell(row=row_idx, column=col)
        if cell.value is not None:
            return True
    return False


def _detect_numeric_cols(ws, row_min, row_max, col_min, col_max):
    """Detecta columnas que contienen numeros (para formatear como numerico)."""
    numeric_cols = set()
    for col in range(col_min, col_max + 1):
        num_count = 0
        total = 0
        for row in range(row_min + 1, min(row_max + 1, row_min + 20)):
            cell = ws.cell(row=row, column=col)
            if cell.value is not None:
                total += 1
                if isinstance(cell.value, (int, float)):
                    num_count += 1
        if total > 0 and num_count / total > 0.6:
            numeric_cols.add(col)
    return numeric_cols


def _detect_header_row(ws):
    """Devuelve el indice de la primera fila que parece cabecera (tiene texto)."""
    for row in ws.iter_rows():
        for cell in row:
            if cell.value and isinstance(cell.value, str) and len(cell.value.strip()) > 0:
                return cell.row
    return 1


def _detect_table_bounds(ws):
    """Detecta los limites de la region de datos activa."""
    min_r = ws.min_row or 1
    max_r = ws.max_row or 1
    min_c = ws.min_column or 1
    max_c = ws.max_column or 1
    return min_r, max_r, min_c, max_c


def _auto_column_width(ws, min_col, max_col, max_row, max_width=50, min_width=8):
    """Ajusta anchos de columna al contenido, con limites."""
    for col in range(min_col, max_col + 1):
        col_letter = get_column_letter(col)
        max_len = min_width
        for row in range(1, min(max_row + 1, 200)):
            cell = ws.cell(row=row, column=col)
            if cell.value is not None:
                try:
                    cell_len = len(str(cell.value))
                    # Celdas combinadas: no sobredimensionar
                    if cell_len > max_len:
                        max_len = cell_len
                except Exception:
                    pass
        # No tocar columnas muy anchas (ya configuradas manualmente)
        current = ws.column_dimensions[col_letter].width
        if current and current > max_width:
            continue   # respetar anchos manuales amplios
        new_width = min(max(max_len + 2, min_width), max_width)
        ws.column_dimensions[col_letter].width = new_width


def style_excel_sheet(ws, sheet_name=''):
    """Aplica la paleta corporativa a una hoja Excel."""
    min_r, max_r, min_c, max_c = _detect_table_bounds(ws)
    if max_r < 1 or max_c < 1:
        return

    header_row = _detect_header_row(ws)
    numeric_cols = _detect_numeric_cols(ws, header_row, max_r, min_c, max_c)

    # Formato numerico espanol: 2 decimales con punto de miles
    NUM_FMT    = '#.##0,00'   # formato ES (Excel interpreta segun locale)
    NUM_FMT_0  = '#.##0'

    for row_idx in range(min_r, max_r + 1):
        is_header = (row_idx == header_row)
        is_alt    = (row_idx % 2 == 0)
        is_total  = False
        # Heuristica de fila total/subtotal
        first_cell = ws.cell(row=row_idx, column=min_c)
        if first_cell.value and isinstance(first_cell.value, str):
            val_upper = first_cell.value.strip().upper()
            if any(kw in val_upper for kw in ('TOTAL', 'SUBTOTAL', 'SUMA', 'IMPORTE')):
                is_total = True

        for col_idx in range(min_c, max_c + 1):
            cell = ws.cell(row=row_idx, column=col_idx)

            # No tocar celdas de formulas en cabecera (raro pero seguro)
            if is_header:
                cell.fill      = _fill(Palette.HEADER_BG)
                cell.font      = _font(bold=True, color=Palette.HEADER_FG)
                cell.alignment = _align('center', 'center', wrap=True)
                cell.border    = _border('thin')
            elif is_total:
                cell.fill      = _fill(Palette.SUBTOTAL_BG)
                cell.font      = _font(bold=True)
                cell.alignment = _align('right' if col_idx in numeric_cols else 'left')
                cell.border    = _border('medium')
            else:
                cell.fill      = _fill(Palette.ALT_ROW if is_alt else Palette.WHITE)
                cell.font      = _font()
                cell.alignment = _align('right' if col_idx in numeric_cols else 'left')
                cell.border    = _border('thin')

            # Formato numerico: solo celdas de datos, no formulas ya formateadas
            if col_idx in numeric_cols and isinstance(cell.value, (int, float)):
                if isinstance(cell.value, float) and cell.value != int(cell.value):
                    cell.number_format = NUM_FMT
                else:
                    cell.number_format = NUM_FMT_0

    # Altura de fila cabecera
    ws.row_dimensions[header_row].height = 22

    # Congelar paneles en la primera fila de datos
    freeze_cell = ws.cell(row=header_row + 1, column=min_c)
    ws.freeze_panes = freeze_cell.coordinate

    # Anchos de columna
    _auto_column_width(ws, min_c, max_c, max_r)

    # Setup de impresion
    ws.print_title_rows = f'${header_row}:${header_row}'
    ws.page_setup.orientation = 'landscape'
    ws.page_setup.fitToPage   = True
    ws.page_setup.fitToWidth  = 1
    ws.page_setup.fitToHeight = 0
    ws.sheet_view.showGridLines = True


def process_excel(path: Path, dry_run=False, backup=True, log=print):
    """Aplica estilos a un archivo Excel."""
    if not HAVE_OPENPYXL:
        log(f'  SKIP {path.name}: openpyxl no disponible')
        return False

    if path.suffix.lower() not in ('.xlsx', '.xlsm'):
        log(f'  SKIP {path.name}: formato no soportado (usar .xlsx)')
        return False

    if path.name.startswith('~'):
        return False  # temporal de Excel abierto

    log(f'  Excel: {path.name}')

    if dry_run:
        log(f'    [DRY-RUN] no se modifica')
        return True

    if backup:
        bak = path.parent / (path.stem + '_ORIG' + path.suffix)
        if not bak.exists():
            shutil.copy2(path, bak)
            log(f'    Backup: {bak.name}')

    try:
        wb = openpyxl.load_workbook(path, data_only=False)
        for ws_name in wb.sheetnames:
            ws = wb[ws_name]
            if ws.sheet_state != 'visible':
                continue
            if ws.max_row and ws.max_row > 0:
                style_excel_sheet(ws, ws_name)
                log(f'    Hoja OK: {ws_name} ({ws.max_row}f x {ws.max_column}c)')
        wb.save(path)
        return True
    except Exception as e:
        log(f'    ERROR: {e}')
        return False


# ══════════════════════════════════════════════════════════════════════════════
# WORD
# ══════════════════════════════════════════════════════════════════════════════

def _set_cell_bg(cell, hex_color):
    """Establece color de fondo de celda Word via XML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    # Eliminar shd existente
    for existing in tcPr.findall(qn('w:shd')):
        tcPr.remove(existing)
    tcPr.append(shd)


def _set_border_side_xml(tc_borders, side, color):
    """Añade un borde lateral a la celda Word."""
    el = OxmlElement(f'w:{side}')
    el.set(qn('w:val'),   'single')
    el.set(qn('w:sz'),    '4')
    el.set(qn('w:space'), '0')
    el.set(qn('w:color'), color)
    tc_borders.append(el)


def _set_cell_borders(cell, color=Palette.BORDER_COLOR):
    """Aplica bordes uniformes a una celda Word."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(existing)
    tc_borders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        _set_border_side_xml(tc_borders, side, color)
    tcPr.append(tc_borders)


def style_word_table(table, is_first_row_header=True):
    """Aplica la paleta corporativa a una tabla Word."""
    for row_idx, row in enumerate(table.rows):
        is_header = (row_idx == 0 and is_first_row_header)
        is_alt    = (row_idx % 2 == 0) and not is_header

        for cell in row.cells:
            # Fondo
            if is_header:
                _set_cell_bg(cell, Palette.HEADER_BG)
            elif is_alt:
                _set_cell_bg(cell, Palette.ALT_ROW)
            else:
                _set_cell_bg(cell, Palette.WHITE)

            # Bordes
            _set_cell_borders(cell)

            # Fuente y alineacion en parrafos
            for para in cell.paragraphs:
                para.paragraph_format.space_before = Pt(2)
                para.paragraph_format.space_after  = Pt(2)
                for run in para.runs:
                    run.font.name = Palette.FONT_NAME
                    run.font.size = Pt(Palette.FONT_SIZE_DOC)
                    if is_header:
                        run.font.bold  = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    else:
                        run.font.bold  = False
                        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

                # Si la celda tiene texto pero no hay runs, aplicar formato al parrafo
                if not para.runs and para.text:
                    # Crear run con el texto existente
                    run = para.add_run()
                    run.font.name = Palette.FONT_NAME
                    run.font.size = Pt(Palette.FONT_SIZE_DOC)
                    if is_header:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def style_word_doc(doc_path: docx.Document, log=print):
    """Aplica estilo a todas las tablas de un documento Word."""
    n_tables = len(doc_path.tables)
    if n_tables == 0:
        log(f'    Sin tablas')
        return 0

    styled = 0
    for i, table in enumerate(doc_path.tables):
        if len(table.rows) < 2:
            continue  # tablas de 1 fila: probablemente diseño, no datos
        try:
            style_word_table(table)
            styled += 1
        except Exception as e:
            log(f'    Tabla {i+1}: ERROR {e}')

    # Fuente global del documento (body text)
    try:
        style = doc_path.styles['Normal']
        style.font.name = Palette.FONT_NAME
        style.font.size = Pt(Palette.FONT_SIZE_DOC)
    except Exception:
        pass

    log(f'    {styled}/{n_tables} tablas estilizadas')
    return styled


def process_word(path: Path, dry_run=False, backup=True, log=print):
    """Aplica estilos a un archivo Word."""
    if not HAVE_DOCX:
        log(f'  SKIP {path.name}: python-docx no disponible')
        return False

    if path.suffix.lower() != '.docx':
        return False

    if path.name.startswith('~') or '_ORIG' in path.name:
        return False

    log(f'  Word:  {path.name}')

    if dry_run:
        log(f'    [DRY-RUN] no se modifica')
        return True

    if backup:
        bak = path.parent / (path.stem + '_ORIG' + path.suffix)
        if not bak.exists():
            shutil.copy2(path, bak)
            log(f'    Backup: {bak.name}')

    try:
        document = docx.Document(str(path))
        n = style_word_doc(document, log=log)
        document.save(str(path))
        return n > 0
    except Exception as e:
        log(f'    ERROR: {e}')
        return False


# ══════════════════════════════════════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════════════════════════════════════

def find_project_root(start=None):
    if start is None:
        start = Path.cwd()
    for p in [Path(start)] + list(Path(start).parents):
        if (p / 'CLAUDE.md').exists() or (p / 'PRESUPUESTO').exists():
            return p
    return Path(start)


def collect_excel_files(base: Path, skip_patterns=None):
    skip_patterns = skip_patterns or ['~$', '_ORIG', 'DONOR', '/FUENTES/']
    files = []
    for f in sorted(base.rglob('*.xlsx')):
        skip = False
        for pat in skip_patterns:
            if pat in f.name or pat in str(f):
                skip = True
                break
        if not skip:
            files.append(f)
    return files


def collect_word_files(base: Path, skip_patterns=None):
    skip_patterns = skip_patterns or ['~$', '_ORIG', 'DONOR', '/Plantillas/']
    files = []
    for f in sorted(base.rglob('*.docx')):
        skip = False
        for pat in skip_patterns:
            if pat in f.name or pat in str(f):
                skip = True
                break
        if not skip:
            files.append(f)
    return files


def main():
    parser = argparse.ArgumentParser(
        description='Estandariza estilo Excel y Word del expediente POU'
    )
    parser.add_argument('--proyecto',   default=None)
    parser.add_argument('--excel-dir',  default=None)
    parser.add_argument('--word-dir',   default=None)
    parser.add_argument('--file',       default=None, help='Procesar un archivo concreto')
    parser.add_argument('--dry-run',    action='store_true')
    parser.add_argument('--excel-only', action='store_true')
    parser.add_argument('--word-only',  action='store_true')
    parser.add_argument('--no-backup',  action='store_true')
    args = parser.parse_args()

    project_root = Path(args.proyecto) if args.proyecto else find_project_root()
    do_excel = not args.word_only
    do_word  = not args.excel_only
    backup   = not args.no_backup
    dry_run  = args.dry_run

    tag = '[DRY-RUN] ' if dry_run else ''

    print(f'\n{"═"*58}')
    print(f'  {tag}ESTANDARIZAR ESTILOS — {project_root.name[:35]}')
    print(f'  {datetime.now().strftime("%Y-%m-%d %H:%M")}')
    print(f'  Paleta: Calibri · Azul institucional #1F3B6B')
    print(f'{"═"*58}\n')

    ok_xl = err_xl = 0
    ok_doc = err_doc = 0

    if args.file:
        p = Path(args.file)
        if p.suffix.lower() in ('.xlsx', '.xlsm') and do_excel:
            r = process_excel(p, dry_run=dry_run, backup=backup)
            if r: ok_xl += 1
            else: err_xl += 1
        elif p.suffix.lower() == '.docx' and do_word:
            r = process_word(p, dry_run=dry_run, backup=backup)
            if r: ok_doc += 1
            else: err_doc += 1
    else:
        anejos_root = Path(args.excel_dir) if args.excel_dir else project_root / 'DOCS - ANEJOS'

        if do_excel:
            print('── Excel ──────────────────────────────────────────────')
            xl_files = collect_excel_files(anejos_root)
            if not xl_files:
                print('  Sin archivos Excel encontrados')
            for f in xl_files:
                r = process_excel(f, dry_run=dry_run, backup=backup)
                if r: ok_xl += 1
                else: err_xl += 1

        if do_word:
            print('\n── Word ───────────────────────────────────────────────')
            word_root = Path(args.word_dir) if args.word_dir else project_root / 'DOCS - ANEJOS'
            doc_files = collect_word_files(word_root)
            if not doc_files:
                print('  Sin archivos Word encontrados')
            for f in doc_files:
                r = process_word(f, dry_run=dry_run, backup=backup)
                if r: ok_doc += 1
                else: err_doc += 1

    print(f'\n{"═"*58}')
    print(f'  Excel: {ok_xl} OK  {err_xl} errores')
    print(f'  Word:  {ok_doc} OK  {err_doc} errores')
    if backup and not dry_run:
        print(f'  Backups _ORIG creados donde no existian')
    print(f'{"═"*58}\n')


if __name__ == '__main__':
    main()
